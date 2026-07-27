"""
docx_generator.py — Word export of the 9-section IC memo (Tâche "Mémo IC et
deck au format IC professionnel").

Produces a professional .docx from the already-generated Markdown memo
(`Deal.ic_memo`) and the shared structured context built by
`ic_context.build_ic_context` — the SAME context object the PPTX deck
generator consumes, so both documents can never show different figures.

This module is a formatter, not a generator of content:
  - Every TABLE (headline figures, financial analysis, sources & uses,
    sensitivity grid, self-check) is built directly from real data
    (Deal / LBOScenario / CompSet), independent of the LLM — it renders
    even if no memo has been generated yet.
  - Every NARRATIVE paragraph is extracted verbatim from the 9 H2 sections
    of `Deal.ic_memo` (never rewritten, never invented here).

Point d'entrée :
    generate_memo_docx(deal, reference_scenario=None, ic_context=None) -> io.BytesIO
"""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

from api.services.ma_engine.ic_context import (
    MEMO_SECTIONS,
    build_ic_context,
    extract_markdown_section,
)

_ACCENT = RGBColor(0x00, 0x66, 0xB3)
_MUTED = RGBColor(0x6C, 0x75, 0x7D)
_DARK = RGBColor(0x1A, 0x1A, 0x2E)
_BAD = RGBColor(0xB1, 0x3B, 0x3B)
_GOOD = RGBColor(0x1F, 0x7A, 0x4D)


def _money(value: float | None) -> str:
    if value is None:
        return "Not available"
    return f"€{value:,.0f}".replace(",", " ")


def _pct(value: float | None, *, already_pct: bool = False) -> str:
    if value is None:
        return "N/A"
    v = value if already_pct else value * 100
    return f"{v:.1f}%"


def _add_bold_aware_paragraph(doc: Document, text: str, *, style: str | None = None):
    """Adds a paragraph, rendering **bold** Markdown runs as real bold runs."""
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def _add_heading(doc: Document, text: str, *, level: int = 2):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = _ACCENT
    return heading


def _render_section_narrative(doc: Document, markdown_text: str | None, heading: str) -> None:
    """Renders the LLM-written narrative for one of the 9 sections, extracted
    verbatim from `Deal.ic_memo` — never rewritten here. Falls back to an
    explicit notice (not a fabricated paragraph) if the memo has not been
    generated yet or the section is absent."""
    lines = extract_markdown_section(markdown_text, heading)
    if not lines:
        p = doc.add_paragraph("IC memo narrative not available for this section — generate the memo first.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _MUTED
        return
    for line in lines:
        if line.startswith(("- ", "* ")):
            _add_bold_aware_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            _add_bold_aware_paragraph(doc, re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            _add_bold_aware_paragraph(doc, line)


def _set_cell_shading(cell, hex_color: str) -> None:
    """Sets a table cell's background fill (no native python-docx API for
    shading — raw OOXML `w:shd`)."""
    shd = cell._tc.get_or_add_tcPr()
    el = shd.makeelement(qn("w:shd"), {qn("w:fill"): hex_color})
    shd.append(el)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1A1A2E")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = str(val)
    doc.add_paragraph()


# ============================================================
# Headline (at-a-glance) table — same headline figures the deck's cover /
# key-figures slides show, so a reader who only skims either document sees
# the same numbers first.
# ============================================================

def _add_at_a_glance_table(doc: Document, deal, ic_context: dict) -> None:
    _add_heading(doc, "At a Glance", level=2)
    fin = ic_context["financials"]
    rows = [
        ["Revenue", fin["revenue_qualified"]],
        ["EBITDA", fin["ebitda_qualified"]],
        ["Enterprise Value", fin["enterprise_value_qualified"]],
    ]
    scenario = ic_context.get("reference_lbo_scenario")
    if scenario:
        rows.append(["Entry Multiple (LBO scenario)", scenario["entry_multiple_qualified"]])
        rows.append(["Leverage at Entry", f"{scenario['leverage_entry_x_ebitda']:.1f}x EBITDA" if scenario.get("leverage_entry_x_ebitda") is not None else "N/A"])
        rows.append(["Projected IRR", _pct(scenario.get("irr")) if scenario.get("irr") is not None else "N/A"])
        rows.append(["Projected MOIC", f"{scenario['moic']:.2f}x" if scenario.get("moic") is not None else "N/A"])
    else:
        rows.append(["LBO Scenario", "No saved scenario for this deal — Deal Terms and Returns sections marked for due diligence."])
    _add_table(doc, ["Metric", "Value"], rows)


# ============================================================
# IV. Financial Analysis — real financial table, independent of the LLM.
# ============================================================

def _add_financial_table(doc: Document, ic_context: dict) -> None:
    fin = ic_context["financials"]
    rows = [
        ["Revenue", fin["revenue_qualified"]],
        ["EBITDA", fin["ebitda_qualified"]],
        ["EBITDA Margin", _pct(fin["ebitda_margin_pct"], already_pct=True) if fin["ebitda_margin_pct"] is not None else "N/A"],
        ["Enterprise Value", fin["enterprise_value_qualified"]],
        ["EV / Revenue", f"{fin['ev_revenue_multiple']:.2f}x" if fin.get("ev_revenue_multiple") else "N/A"],
        ["EV / EBITDA (deal-level, implicit)", f"{fin['ev_ebitda_multiple']:.2f}x" if fin.get("ev_ebitda_multiple") else "N/A"],
        ["Quality of Earnings", fin["quality_of_earnings"]],
        ["Working Capital", fin["working_capital"]],
        ["Capex Detail", fin["capex_detail"]],
    ]
    _add_table(doc, ["Metric", "Value"], rows)


# ============================================================
# VI. Deal Terms & Structure — Sources & Uses table, real LBO scenario data.
# ============================================================

def _add_sources_and_uses_table(doc: Document, ic_context: dict) -> None:
    scenario = ic_context.get("reference_lbo_scenario")
    sau = ic_context.get("sources_and_uses")
    if not scenario or not sau:
        p = doc.add_paragraph(
            "No LBO scenario has been modelled for this deal — deal terms, capital "
            "structure and leverage are marked for due diligence."
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _MUTED
        return

    rows = [["Entry Multiple", scenario["entry_multiple_qualified"]]]
    # Tâche "P1 : physique financière du modèle LBO" (Partie E) — Uses n'est
    # plus l'Enterprise Value seule : frais de transaction/financement et
    # cash minimum financé à la clôture s'y ajoutent.
    rows.append(["USES — Enterprise Value (Purchase of Target)", _money(sau.get("entry_ev"))])
    rows.append(["— Transaction Fees (advisory/legal/DD)", _money(sau.get("entry_transaction_fees"))])
    rows.append(["— Financing Fees (debt arrangement)", _money(sau.get("entry_financing_fees"))])
    rows.append(["— Minimum Cash (funded at closing)", _money(sau.get("entry_min_cash"))])
    rows.append(["Total Uses", _money(sau.get("entry_uses_total"))])
    tranches = sau.get("tranches")
    if tranches:
        for t in tranches:
            label = f"SOURCES — {t.get('name', 'Debt Tranche')} ({t.get('amortization', 'bullet')}, {_pct(t.get('interest_rate'))})"
            rows.append([label, _money(t.get("amount"))])
    else:
        rows.append(["SOURCES — Senior Debt (single tranche)", _money(sau.get("entry_debt"))])
    rows.append(["— Equity (sponsor, plug)", _money(sau.get("entry_equity"))])
    rows.append(["Leverage at Entry", f"{sau['leverage_entry']:.1f}x EBITDA" if sau.get("leverage_entry") is not None else "N/A"])
    _add_table(doc, ["Item", "Amount"], rows)

    # Tâche "P2 : crédibilité de la thèse" (Partie A) — sous le seuil de CA,
    # ce LBO standalone est indicatif : la même mise en garde que le deck,
    # ouvrant la section, jamais reléguée à une note de bas de page discrète.
    sizing = scenario.get("sizing_guidance") or {}
    if sizing.get("is_indicative") and sizing.get("note"):
        p = doc.add_paragraph(sizing["note"])
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = _BAD
        doc.add_paragraph()

    if scenario.get("ebitda_reconciliation_note"):
        p = doc.add_paragraph(scenario["ebitda_reconciliation_note"])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = _MUTED
    if scenario.get("valuation_reconciliation_note"):
        p = doc.add_paragraph(scenario["valuation_reconciliation_note"])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = _MUTED
        doc.add_paragraph()


# ============================================================
# VII. Returns Analysis — real returns KPIs + sensitivity grid.
# ============================================================

def _add_returns_table(doc: Document, ic_context: dict) -> None:
    scenario = ic_context.get("reference_lbo_scenario")
    if not scenario:
        p = doc.add_paragraph(
            "No LBO scenario has been modelled for this deal — returns are marked "
            "for due diligence."
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _MUTED
        return

    rows = [
        ["Holding Period", f"{int(scenario['holding_period_years'])} years" if scenario.get("holding_period_years") else "N/A"],
        ["Exit Multiple", f"{scenario['exit_multiple']:.2f}x" if scenario.get("exit_multiple") is not None else "N/A"],
        ["Projected IRR", _pct(scenario.get("irr")) if scenario.get("irr") is not None else "N/A"],
        ["Projected MOIC", f"{scenario['moic']:.2f}x" if scenario.get("moic") is not None else "N/A"],
        ["Revenue Growth (p.a.)", _pct(scenario.get("revenue_growth_pct"), already_pct=True) if scenario.get("revenue_growth_pct") is not None else "N/A"],
        ["Interest Rate (weighted avg.)", _pct(scenario.get("interest_rate_pct"), already_pct=True) if scenario.get("interest_rate_pct") is not None else "N/A"],
    ]
    _add_table(doc, ["Assumption / Result", "Value"], rows)

    # Tâche "P2 : crédibilité de la thèse" (Partie B) — "un mémo d'IC
    # mono-scénario n'existe pas" : base ET downside côte à côte, jamais le
    # base case seul quand un downside a été généré pour ce deal.
    downside = ic_context.get("downside_scenario")
    doc.add_paragraph()
    p = doc.add_paragraph("Base Case vs Downside Case")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = _ACCENT
    if downside:
        haircut = downside.get("revenue_haircut_pct")
        delta = downside.get("exit_multiple_delta")
        compare_rows = [
            [
                "Key assumption vs base",
                "—",
                (
                    f"Revenue {haircut * -100:.0f}%" if haircut is not None else "—"
                ) + (
                    f", exit multiple {delta:+.1f}x" if delta is not None else ""
                ),
            ],
            ["Exit Multiple", f"{scenario['exit_multiple']:.2f}x" if scenario.get("exit_multiple") is not None else "N/A",
             f"{downside['exit_multiple']:.2f}x" if downside.get("exit_multiple") is not None else "N/A"],
            ["Projected IRR", _pct(scenario.get("irr")) if scenario.get("irr") is not None else "N/A",
             _pct(downside.get("irr")) if downside.get("irr") is not None else "N/A"],
            ["Projected MOIC", f"{scenario['moic']:.2f}x" if scenario.get("moic") is not None else "N/A",
             f"{downside['moic']:.2f}x" if downside.get("moic") is not None else "N/A"],
        ]
        _add_table(doc, ["Metric", "Base Case", downside.get("label", "Downside Case")], compare_rows)
    else:
        p = doc.add_paragraph("Only a base case has been modelled for this deal — no downside scenario available.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _MUTED
    doc.add_paragraph()

    sensitivity = ic_context.get("sensitivity")
    if not sensitivity:
        p = doc.add_paragraph(
            "Sensitivity analysis not available (requires a saved LBO scenario with a "
            "complete sector profile) — marked for due diligence."
        )
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = _MUTED
        return

    doc.add_paragraph()
    p = doc.add_paragraph("Sensitivity — IRR / MOIC by Exit Multiple × Leverage")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = _ACCENT

    headers = ["Leverage \\ Exit Multiple"] + [f"{ex:.1f}x" for ex in sensitivity["exit_axis"]]
    rows = []
    for row in sensitivity["grid"]:
        cells = [f"{row['leverage']:.1f}x"]
        for cell in row["cells"]:
            irr_txt = _pct(cell["irr"]) if cell["irr"] is not None else "N/A"
            moic_txt = f"{cell['moic']:.2f}x" if cell["moic"] is not None else "N/A"
            cells.append(f"{irr_txt} / {moic_txt}")
        rows.append(cells)
    _add_table(doc, headers, rows)

    p = doc.add_paragraph(sensitivity["method"])
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9.5)
    p.runs[0].font.color.rgb = _MUTED
    doc.add_paragraph()


# ============================================================
# VIII. Risk Factors — self-check ALWAYS rendered as a table, independent of
# whether the LLM narrative also mentioned it (belt and suspenders: a known
# data limitation must never depend on the model's compliance to surface).
# ============================================================

def _add_self_check_table(doc: Document, ic_context: dict) -> None:
    self_check = ic_context.get("self_check") or {}
    checks = self_check.get("checks") or []
    if not checks:
        return

    doc.add_paragraph()
    p = doc.add_paragraph("Data & Model Consistency Self-Check")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = _ACCENT

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Check", "Result", "Detail"]):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "1A1A2E")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for check in checks:
        cells = table.add_row().cells
        cells[0].text = check["name"]
        cells[1].text = "Passed" if check["passed"] else "Discrepancy found"
        cells[1].paragraphs[0].runs[0].font.color.rgb = _GOOD if check["passed"] else _BAD
        cells[1].paragraphs[0].runs[0].bold = True
        cells[2].text = check["detail"]
    doc.add_paragraph()


# ============================================================
# Point d'entrée
# ============================================================

def generate_memo_docx(deal, reference_scenario=None, ic_context: dict | None = None) -> io.BytesIO:
    """Generates the 9-section IC memo as a Word (.docx) document.

    Args:
        deal: `Deal` instance (should have `ic_memo` already generated —
            tables render regardless, narrative shows a notice if absent).
        reference_scenario: optional `LBOScenario` — used only if
            `ic_context` is not supplied (rebuilds it internally).
        ic_context: the structured context from
            `ic_context.build_ic_context` — pass the SAME object used to
            generate the PPTX deck to guarantee the two documents agree.

    Returns:
        io.BytesIO ready to stream (.docx).
    """
    if ic_context is None:
        ic_context = build_ic_context(deal, reference_scenario, None)

    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # --- Header ---
    title = doc.add_heading("Investment Committee Memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = _DARK

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(deal.target_name or "Unnamed Target")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = _ACCENT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Acquirer: {deal.acquirer_name or 'N/A'}  •  Sector: {deal.sector or 'N/A'}  •  "
        f"Country: {deal.country or 'N/A'}  •  Generated on {datetime.now().strftime('%d %B %Y')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = _MUTED

    doc.add_paragraph()

    _add_at_a_glance_table(doc, deal, ic_context)

    memo = deal.ic_memo

    # --- I. Executive Summary ---
    _add_heading(doc, MEMO_SECTIONS[0])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[0])

    # --- II. Company Overview ---
    _add_heading(doc, MEMO_SECTIONS[1])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[1])

    # --- III. Industry & Market ---
    _add_heading(doc, MEMO_SECTIONS[2])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[2])

    # --- IV. Financial Analysis ---
    _add_heading(doc, MEMO_SECTIONS[3])
    _add_financial_table(doc, ic_context)
    _render_section_narrative(doc, memo, MEMO_SECTIONS[3])

    # --- V. Investment Thesis ---
    _add_heading(doc, MEMO_SECTIONS[4])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[4])

    # --- VI. Deal Terms & Structure ---
    _add_heading(doc, MEMO_SECTIONS[5])
    _add_sources_and_uses_table(doc, ic_context)
    _render_section_narrative(doc, memo, MEMO_SECTIONS[5])

    # --- VII. Returns Analysis ---
    _add_heading(doc, MEMO_SECTIONS[6])
    _add_returns_table(doc, ic_context)
    _render_section_narrative(doc, memo, MEMO_SECTIONS[6])

    # --- VIII. Risk Factors ---
    _add_heading(doc, MEMO_SECTIONS[7])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[7])
    _add_self_check_table(doc, ic_context)

    # --- IX. Recommendation ---
    _add_heading(doc, MEMO_SECTIONS[8])
    _render_section_narrative(doc, memo, MEMO_SECTIONS[8])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
